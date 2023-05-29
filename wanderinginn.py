from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait as W
from selenium.webdriver.support import expected_conditions as E
import time
import docx
url="https://wanderinginn.com/"
# Make sure to upgrade your chrome driver to the present version of whatever your chrome is 
# To check the present version of chrome go to a new tab in chrome and type chrome://version
# After that make sure to go to https://chromedriver.chromium.org/downloads to download chromedriver version similar to your chrome
PATH = "C:\\drivers\\chromedriver.exe"
a=[]
b=[]
wait_time_out =15
driver = webdriver.Chrome(PATH)
driver.get(url)
wait_variable = W(driver,wait_time_out)

time.sleep(5)
driver.find_element("link text","Table of Contents").click()
links = wait_variable.until(E.visibility_of_any_elements_located((By.TAG_NAME,"a")))
time.sleep(4)

doc = docx.Document()
for link in links:
    x=link.text
    if x=='Blog at WordPress.com.':
        break
    a.append(x)
y=a.index('1.00')
print(a[y:])
b=a[y:].copy()

for i in b:
    print(i)
    time.sleep(2)
    driver.find_element("link text","Table of Contents").click()
    driver.find_element("link text",i).click()
    summary=driver.find_element("css selector",'.entry-content').text
    doc.add_heading(i, 0)
    doc.add_paragraph(summary)
    doc.save('wanderinginn.docx')
    time.sleep(2)
driver.close()
